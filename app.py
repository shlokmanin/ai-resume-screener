from flask import Flask, render_template, request, flash, redirect, url_for
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import logging
from datetime import datetime
import os
from dotenv import load_dotenv
from models import db, User, Resume, Feedback
import PyPDF2
from flask_wtf.csrf import CSRFProtect
import re
import nltk
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from flask import send_file
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from io import BytesIO

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-secret-key-do-not-use-in-production')
app.config['UPLOAD_FOLDER'] = os.getenv('UPLOAD_FOLDER', 'uploads')
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('MAX_CONTENT_LENGTH', 5*1024*1024))  # 5MB

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Database config
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///database.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db.init_app(app)

# CSRF Protection
csrf = CSRFProtect(app)
app.config['WTF_CSRF_ENABLED'] = True

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        
        if password != confirm_password:
            flash('Passwords do not match', 'error')
            return render_template('signup.html')
        
        if User.query.filter_by(email=email).first():
            flash('Email already registered', 'error')
            return render_template('signup.html')
        
        user = User(username=username, email=email)
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        flash('Account created successfully! Please login.', 'success')
        return redirect(url_for('login'))
    
    return render_template('signup.html')

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        user = User.query.filter_by(email=email).first()
        if user and user.check_password(password):
            login_user(user)
            flash('Login successful!', 'success')
            return redirect(url_for('dashboard'))
        flash('Invalid email or password', 'error')
    return render_template('login.html')

@login_required
@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'POST':
        if 'resume' not in request.files:
            flash('No file selected', 'error')
            return redirect(request.url)
        
        file = request.files['resume']
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)
        
        if file and file.filename.lower().endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Extract text from PDF
            extracted_text = ""
            with open(filepath, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    extracted_text += page.extract_text() + "\n"
            
            # Download nltk data if not present
            try:
                nltk.data.find('tokenizers/punkt')
            except LookupError:
                nltk.download('punkt')
            try:
                nltk.data.find('corpora/stopwords')
            except LookupError:
                nltk.download('stopwords')
            
            # Extract skills
            skills_keywords = [
                'python', 'java', 'sql', 'javascript', 'html', 'css', 'react', 'node.js', 
                'machine learning', 'data science', 'deep learning', 'nlp', 'computer vision',
                'aws', 'docker', 'kubernetes', 'git', 'django', 'flask', 'tensorflow', 'pytorch'
            ]
            
            # Simple skill extraction
            text_lower = extracted_text.lower()
            detected_skills = [skill for skill in skills_keywords if skill in text_lower]
            
            # Save to database
            new_resume = Resume(
                user_id=current_user.id,
                file_name=filename,
                extracted_text=extracted_text,
                skills=', '.join(detected_skills),
                date_uploaded=datetime.utcnow()
            )
            db.session.add(new_resume)
            db.session.commit()
            
            flash(f'Resume processed! Found {len(detected_skills)} skills', 'success')
            return redirect(url_for('results', resume_id=new_resume.id))
        
        flash('Please upload PDF only', 'error')
    
    return render_template('upload.html')

@login_required
@app.route('/dashboard')
def dashboard():
    return render_template('dashboard.html')

@login_required
@app.route('/results')
def results():
    # Get latest resume for user
    latest_resume = Resume.query.filter_by(user_id=current_user.id).order_by(Resume.id.desc()).first()
    if not latest_resume:
        flash('No resume found. Please upload first.', 'error')
        return redirect(url_for('upload'))
    return render_template('results.html', resume=latest_resume)

@login_required
@app.route('/results/<int:resume_id>', methods=['GET', 'POST'])
def results_detail(resume_id):
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first_or_404()
    
    if request.method == 'POST':
        job_desc = request.form['job_desc']
        
        resume_text = resume.extracted_text or ""
        job_desc = request.form.get("job_desc", "")
        
        if resume_text and job_desc:
            # Keyword Matching
            job_words = set(job_desc.lower().split())
            resume_words = set(resume_text.lower().split())
            matched_words = job_words.intersection(resume_words)
            keyword_score = (len(matched_words) / len(job_words)) if len(job_words) > 0 else 0
            
            # TF-IDF Similarity
            documents = [resume_text, job_desc]
            vectorizer = TfidfVectorizer(stop_words='english')
            tfidf_matrix = vectorizer.fit_transform(documents)
            tfidf_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            
            # Final Combined Score
            final_score = (0.7 * keyword_score + 0.3 * tfidf_score) * 100
            match_score = round(final_score, 2)
            
# Advanced AI Recommendations (limit to 7)
            recommendations = []
            
            # 1. Missing Skills (top 5)
            job_keywords = [word for word in job_desc.lower().split() if len(word) > 3 and word.isalpha()]
            resume_lower = resume_text.lower()
            missing_skills = [skill for skill in set(job_keywords)[:10] if skill not in resume_lower][:5]
            for skill in missing_skills:
                recommendations.append(f"📌 Add '{skill.title()}' skill to improve match")
            
            # 2. Resume Improvement Tips
            text_length = len(resume_text.split())
            if text_length < 200:
                recommendations.append("📄 Expand content: Add detailed project descriptions (aim for 300+ words)")
            if 'project' not in resume_lower and 'projects' not in resume_lower:
                recommendations.append("🚀 Add 'Projects' section with technologies used")
            if 'experience' not in resume_lower and 'internship' not in resume_lower:
                recommendations.append("💼 Include internship/work experience or academic projects")
            if 'certification' not in resume_lower and 'certificate' not in resume_lower:
                recommendations.append("🏆 Add certifications (Coursera, AWS, Google, etc.)")
            
            # 3. Content Quality Tips
            recommendations.append("✍️ Use action verbs: Developed, Built, Designed, Led, Optimized")
            recommendations.append("📊 Add measurable achievements: 'Improved performance by 30%', 'Reduced load time by 40%'")
            
            # 4. Headline Suggestion
            if resume.skills:
                headline = f"{current_user.username.split()[0]} | {resume.skills.split(',')[0].title()} Expert | Final Year Student"
            else:
                headline = f"{current_user.username.split()[0]} | Aspiring Software Developer"
            recommendations.append(f"🎯 Suggested Headline: \"{headline}\"")
            
            # Limit to 7
            recommendations = recommendations[:7]
            
        else:
            match_score = 0
            recommendations = ["Please enter a job description to get accurate score and recommendations."]
        
        match_score = round(match_score, 2)
        
        resume.match_score = match_score / 100.0  # Store as 0-1 fraction
        resume.job_description = job_desc
        db.session.commit()
        
        flash(f'Match score updated: {match_score}%', 'success')
        return render_template('results.html', resume=resume, match_score=match_score, recommendations=recommendations)
    
    return render_template('results.html', resume=resume)

@login_required
@app.route('/feedback', methods=['GET', 'POST'])
def feedback():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        message = request.form['message']
        rating = int(request.form['rating'])
        
        new_feedback = Feedback(
            name=name,
            email=email,
            message=message,
            rating=rating
        )
        db.session.add(new_feedback)
        db.session.commit()
        flash('Thank you for your feedback!', 'success')
        return redirect(url_for('dashboard'))
    
    return render_template('feedback.html')

def generate_improved_resume(resume, username):
    """Generate structured improved resume text"""
    skills = resume.skills or "Python, SQL, Machine Learning"
    text_lower = resume.extracted_text.lower() if resume.extracted_text else ""
    
    headline = f"{username.split()[0]} | {skills.split(',')[0].title()} Developer | Final Year B.Tech Student"
    
    projects = []
    if 'project' in text_lower:
        projects.append("- Developed ML model for resume screening (Python, scikit-learn)")
        projects.append("- Built full-stack web app with Flask and React")
    else:
        projects.append("- Add your projects with tech stack and achievements")
    
    experience = "Final Year B.Tech Student | Relevant Coursework: ML, Data Structures, DBMS"
    if 'internship' in text_lower:
        experience = "Internship Experience | Academic Projects"
    
    certifications = ["Coursera: Machine Learning (Recommended)", "AWS Cloud Practitioner (Suggested)"]
    
    improved_resume = f"""{username.title()}
{headline}

SKILLS
{skills.upper()}

PROJECTS
""" + "\n".join(projects) + f"""

EXPERIENCE
{experience}

CERTIFICATIONS
""" + "\n".join(certifications) + """

EDUCATION
B.Tech Computer Science | Final Year | GPA: X.XX

Add more sections, use action verbs, quantify achievements!
"""
    return improved_resume

@login_required
@app.route('/generate_resume/<int:resume_id>')
def generate_resume(resume_id):
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first_or_404()
    generated_text = generate_improved_resume(resume, current_user.username)
    return render_template('generate_resume.html', resume=resume, generated_text=generated_text)

@login_required
@app.route('/download/pdf/<int:resume_id>', methods=['POST'])
def download_pdf(resume_id):
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first_or_404()
    edited_text = request.form.get('edited_resume', '')
    if not edited_text:
        edited_text = generate_improved_resume(resume, current_user.username)
    
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 50
    lines = edited_text.split('\n')
    for line in lines:
        if y < 50:
            p.showPage()
            p.setFont("Helvetica", 10)
            y = height - 50
        p.drawString(50, y, line[:80])  # Truncate long lines
        y -= 15
    p.save()
    buffer.seek(0)
    
    from flask import make_response
    response = make_response(send_file(buffer, as_attachment=True, download_name=f"AI_Resume_{resume.file_name.replace('.pdf','')}.pdf", mimetype='application/pdf'))
    return response

@login_required
@app.route('/download/docx/<int:resume_id>', methods=['POST'])
def download_docx(resume_id):
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first_or_404()
    edited_text = request.form.get('edited_resume', '')
    if not edited_text:
        edited_text = generate_improved_resume(resume, current_user.username)
    
    doc = Document()
    doc.add_heading(current_user.username.title(), 0)
    doc.add_heading('Generated by AI Resume Enhancer', level=2)
    doc.add_paragraph(edited_text)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"AI_Resume_{resume.file_name.replace('.pdf','')}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Logged out successfully', 'success')
    return redirect(url_for('index'))

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('500.html'), 500

import os

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        print("Database tables created successfully!")
    print("AI Resume Screening System starting...")
    print(f"Upload folder: {app.config['UPLOAD_FOLDER']}")
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)

